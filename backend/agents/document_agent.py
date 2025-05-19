
# import logging
# import os
# import shutil
# from typing import List, Dict, Any
# from config import get_vectorstore, DOC_DIRS, VECTOR_DIRS
# from docx.table import Table # Import the Table class
# # --- Import cần thiết ---
# # Cần cài đặt: pip install python-docx langchain langchain-openai langchain-chroma python-dotenv langchain_experimental
# from dotenv import load_dotenv
# from docx import Document as DocxDocument
# from langchain_core.documents import Document
# # THAY ĐỔI: Import SemanticChunker từ langchain_experimental
# from langchain_experimental.text_splitter import SemanticChunker
# # from langchain.text_splitter import RecursiveCharacterTextSplitter # Không dùng nữa
# # from langchain_community.vectorstores import Chroma
# from langchain_chroma import Chroma
# from langchain_openai import OpenAIEmbeddings

# # --- Import các hàm tùy chỉnh (Giả định đã được sửa đổi) ---
# from agents.table_agent import extract_tables_from_docx # Cần trả về Markdown table trong page_content
# from agents.ocr_agent import extract_images_and_ocr # Cần trả về OCR text trong page_content

# # --- Cấu hình Logging ---
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
# logger = logging.getLogger(__name__)

# # --- Cấu hình (Nên chuyển vào config.py hoặc .env) ---
# # Giả sử DOC_DIRS được định nghĩa:
# # DOC_DIRS = {"general": "data/word_docs"} # Ví dụ nếu chỉ có 1 thư mục chung
# # VECTOR_STORE_BASE_DIR = "data/vector_stores_semantic" # Đổi tên thư mục để tránh ghi đè
# # CHÚ Ý: Semantic Chunker không dùng chunk_size/overlap trực tiếp
# # Các tham số của nó là embedding model và breakpoint threshold

# # --- Load API Key ---
# load_dotenv()

# # --- Khởi tạo Embedding Function (Cần cho SemanticChunker) ---
# try:
#     # Semantic Chunker yêu cầu embedding function
#     embedding_function = OpenAIEmbeddings()
#     logger.info("✅ OpenAI Embeddings initialized successfully.")
# except Exception as e:
#     logger.error(f"❌ Cannot initialize OpenAI Embeddings: {e}", exc_info=True)
#     exit()

# # --- Helper Functions ---

# def table_to_markdown(table: Table) -> str:
#     """Chuyển đổi bảng Docx thành Markdown (Ví dụ cơ bản)."""
#     # (Giữ nguyên hàm này như phiên bản trước)
#     markdown = ""
#     try:
#         headers = [cell.text.strip().replace("|", "\\|") for cell in table.rows[0].cells]
#         markdown += "| " + " | ".join(headers) + " |\n"
#         markdown += "| " + " | ".join(["---"] * len(headers)) + " |\n"
#         for row in table.rows[1:]:
#             cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
#             markdown += "| " + " | ".join(cells) + " |\n"
#     except Exception as e:
#         logger.warning(f"⚠️ Error converting table to markdown: {e}")
#         markdown = "\n".join(cell.text for row in table.rows for cell in row.cells)
#     return markdown

# # --- Core Processing Functions ---

# def extract_content_from_docx(filepath: str, filename: str, domain: str) -> List[Document]:
#     """Đọc file .docx, trích xuất nội dung (Vẫn giữ nguyên hàm này)."""
#     # (Giữ nguyên logic của hàm này như phiên bản trước)
#     # ... (Code trích xuất paragraph, gọi table_to_markdown, gọi extract_images_and_ocr) ...
#     raw_docs_for_file: List[Document] = []
#     logger.info(f"📄 Processing file: {filename}")
#     try:
#         docx = DocxDocument(filepath)
#         base_metadata = {"source": filename, "domain": domain}
#         para_index = 0
#         for p in docx.paragraphs:
#             text = p.text.strip()
#             if text:
#                 metadata = base_metadata.copy()
#                 metadata.update({"type": "paragraph", "index": f"p_{para_index}"})
#                 raw_docs_for_file.append(Document(page_content=text, metadata=metadata))
#                 para_index += 1

#         table_index = 0
#         for table in docx.tables:
#             markdown_table = table_to_markdown(table)
#             if markdown_table:
#                 metadata = base_metadata.copy()
#                 metadata.update({"type": "table", "index": f"t_{table_index}"})
#                 raw_docs_for_file.append(Document(page_content=markdown_table, metadata=metadata))
#                 table_index += 1
#         try:
#             ocr_docs = extract_images_and_ocr(docx, filename)
#             for i, od in enumerate(ocr_docs):
#                 ocr_metadata = base_metadata.copy()
#                 ocr_metadata.update(od.metadata)
#                 ocr_metadata.update({"type": "ocr_image", "index": f"ocr_{i}"})
#                 od.metadata = ocr_metadata
#                 od.page_content = f"[Nội dung OCR từ ảnh]:\n{od.page_content}"
#                 raw_docs_for_file.append(od)
#         except Exception as e:
#             logger.error(f"❌ Error during OCR in {filename}: {e}", exc_info=False)

#     except Exception as e:
#         logger.error(f"❌ Critical error reading file {filepath}: {e}", exc_info=True)
#         return []
#     logger.info(f"  Extracted {len(raw_docs_for_file)} raw elements from {filename}")
#     return raw_docs_for_file


# # --- THAY ĐỔI CHÍNH Ở HÀM NÀY ---
# def chunk_documents_semantic(docs: List[Document], embedding_func) -> List[Document]:
#     """Chia nhỏ danh sách Document bằng SemanticChunker."""
#     if not docs:
#         return []

#     logger.info(f"⚡ Chunking {len(docs)} raw documents using Semantic Chunker...")
#     # Khởi tạo SemanticChunker với embedding function
#     # Cần thử nghiệm các loại breakpoint_threshold_type:
#     # "percentile" (mặc định, thường dùng 95), "standard_deviation", "interquartile", "gradient"
#     splitter = SemanticChunker(
#         embeddings=embedding_func,
#         breakpoint_threshold_type="percentile" # Bắt đầu với percentile
#         # breakpoint_threshold_amount=0.95 # Có thể thử thay đổi giá trị này
#     )

#     try:
#         # Thực hiện chia nhỏ
#         # Lưu ý: Bước này sẽ gọi API embedding nhiều lần và tốn thời gian/chi phí
#         final_docs = splitter.split_documents(docs)
#         logger.info(f"  Finished semantic chunking. Result: {len(final_docs)} final chunks.")
#         # Không cần bước deduplicate vì semantic chunker tự tạo ra các chunk khác biệt về ngữ nghĩa
#         return final_docs
#     except Exception as e:
#          logger.error(f"❌ Lỗi trong quá trình semantic chunking: {e}", exc_info=True)
#          return []

# def initialize_or_load_vector_store(persist_dir: str, embedding_func, domain: str) -> Chroma:
#     """Khởi tạo hoặc tải Chroma vector store (Vẫn giữ nguyên hàm này)."""
#     # (Giữ nguyên logic của hàm này như phiên bản trước)
#     if not os.path.exists(persist_dir):
#         logger.info(f"✨ Vector store not found at {persist_dir}. Will create a new one.")
#         os.makedirs(os.path.dirname(persist_dir), exist_ok=True)
#         return Chroma(persist_directory=persist_dir, embedding_function=embedding_func, collection_name=domain)
#     else:
#         logger.info(f"🌀 Loading existing vector store from {persist_dir}")
#         try:
#             return Chroma(persist_directory=persist_dir, embedding_function=embedding_func)
#         except Exception as e:
#             logger.error(f"❌ Failed to load vector store from {persist_dir}: {e}. ", exc_info=True)
#             return None

# # --- Main Processing Function ---

# def run_document_processing():
#     """Hàm chính điều phối xử lý (Cập nhật để gọi hàm chunking mới)."""
#     logger.info("🚀 Starting document processing workflow (using Semantic Chunking)...")

#     if not DOC_DIRS:
#         logger.warning("⚠️ DOC_DIRS is empty. No documents to process.")
#         return

#     all_processed_files = 0
#     all_added_chunks = 0

#     # Khởi tạo embedding function một lần ở đây để truyền vào các hàm cần thiết
#     # (Đã khởi tạo ở trên cùng)
#     global embedding_function
#     if embedding_function is None:
#         logger.error("❌ Embedding function not initialized. Exiting.")
#         return

#     for domain, folder_path in DOC_DIRS.items():
#         logger.info(f"📂 Processing domain: '{domain}' from folder: {folder_path}")

#         if not os.path.isdir(folder_path):
#             logger.warning(f"⚠️ Directory not found: {folder_path}. Skipping domain '{domain}'.")
#             continue

#         # persist_dir = os.path.join(VECTOR_DIRS, domain)
#         persist_dir = VECTOR_DIRS[domain]
#         vector_store = initialize_or_load_vector_store(persist_dir, embedding_function, domain)
#         # persist_dir = os.path.join(domain)
#         # vector_store = initialize_or_load_vector_store(persist_dir, embedding_function)

#         if vector_store is None:
#             logger.error(f"❌ Cannot proceed with domain '{domain}' due to vector store error.")
#             continue

#         domain_raw_docs: List[Document] = []
#         processed_files_count = 0
#         for filename in os.listdir(folder_path):
#             if filename.endswith(".docx") and not filename.startswith("~"):
#                 filepath = os.path.join(folder_path, filename)
#                 file_docs = extract_content_from_docx(filepath, filename, domain)
#                 if file_docs:
#                     domain_raw_docs.extend(file_docs)
#                     processed_files_count += 1

#         if not domain_raw_docs:
#             logger.info(f"ℹ️ No processable .docx documents found or extracted from {folder_path}.")
#             continue

#         # --- THAY ĐỔI CHÍNH: Gọi hàm chunking mới ---
#         final_chunks = chunk_documents_semantic(domain_raw_docs, embedding_function)

#         if final_chunks:
#             try:
#                 logger.warning(f"⏳ Adding/Updating {len(final_chunks)} chunks for domain '{domain}' using Semantic Chunker (This may take time and cost)...")
#                 vector_store.add_documents(final_chunks)
#                 # vectorstore.add_documents(final_chunks)
#                 # vector_store.persist()  # ✅ Bắt buộc để lưu file index

#                 all_added_chunks += len(final_chunks)
#                 logger.info(f"✅ Successfully added/updated chunks for domain '{domain}'.")
#             except Exception as e:
#                  logger.error(f"❌ Error adding documents to vector store for domain '{domain}': {e}", exc_info=True)
#         else:
#             logger.info(f"ℹ️ No final semantic chunks to add for domain '{domain}'.")

#         all_processed_files += processed_files_count

#     logger.info("-" * 30)
#     logger.info("🏁 Document processing workflow finished.")
#     logger.info(f"Processed {all_processed_files} files.")
#     logger.info(f"Added/Updated approx. {all_added_chunks} semantic chunks across all domains.")
#     logger.info(f"Vector stores located in: {VECTOR_DIRS}")

# # --- Main Execution Block ---
# if __name__ == "__main__":
#     logger.warning("⚠️ Running processing with Semantic Chunker. This might be slow and consume API credits for embeddings.")
#     run_document_processing()

import os
import shutil
import logging
import subprocess
from typing import List
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_experimental.text_splitter import SemanticChunker
from docx import Document as DocxDocument
from docx.table import Table
from docling.document_converter import DocumentConverter
from langchain_text_splitters import MarkdownHeaderTextSplitter
from dotenv import load_dotenv

# Load .env
load_dotenv()

# Logging setup
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Constants
PUBLIC_FILES_DIR = "frontend/public/files"
DOC_DIRS = {
    "dao_tao": "data/dao_tao",
    "ctsv": "data/ctsv"
}
VECTOR_DIRS = {
    "dao_tao": "data/vector_dao_tao",
    "ctsv": "data/vector_ctsv"
}
embedding_function = OpenAIEmbeddings()

def convert_docx_to_pdf(docx_path: str, output_dir: str):
    try:
        os.makedirs(output_dir, exist_ok=True)
        subprocess.run([
            "soffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            docx_path
        ], check=True)
        logger.info(f"✅ Converted to PDF: {docx_path}")
    except Exception as e:
        logger.error(f"❌ Failed to convert {docx_path} to PDF: {e}")

def table_to_markdown(table: Table) -> str:
    try:
        headers = [cell.text.strip().replace("|", "\\|") for cell in table.rows[0].cells]
        markdown = "| " + " | ".join(headers) + " |\n"
        markdown += "| " + " | ".join(["---"] * len(headers)) + " |\n"
        for row in table.rows[1:]:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            markdown += "| " + " | ".join(cells) + " |\n"
        return markdown
    except Exception as e:
        logger.warning(f"⚠️ Table markdown error: {e}")
        return ""

def extract_from_docx(filepath: str, domain: str) -> List[Document]:
    documents = []
    try:
        doc = DocxDocument(filepath)
        filename = os.path.basename(filepath)
        base_metadata = {"source": filename, "domain": domain}
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                documents.append(Document(page_content=p.text.strip(), metadata={**base_metadata, "type": "paragraph", "index": f"p{i}"}))
        for j, table in enumerate(doc.tables):
            md_table = table_to_markdown(table)
            if md_table:
                documents.append(Document(page_content=md_table, metadata={**base_metadata, "type": "table", "index": f"t{j}"}))
        convert_docx_to_pdf(filepath, PUBLIC_FILES_DIR)
    except Exception as e:
        logger.error(f"❌ Error processing DOCX: {e}")
    return documents

def extract_from_pdf(filepath: str, domain: str) -> List[Document]:
    documents = []
    try:
        converter = DocumentConverter()
        markdown = converter.convert(filepath).document.export_to_markdown()
        splitter = MarkdownHeaderTextSplitter(headers_to_split_on=[])
        docs = splitter.split_text(markdown)
        filename = os.path.basename(filepath)
        for i, doc in enumerate(docs):
            documents.append(Document(page_content=doc.page_content.strip(), metadata={"source": filename, "domain": domain, "index": f"md_{i}"}))
        shutil.copy(filepath, os.path.join(PUBLIC_FILES_DIR, os.path.basename(filepath)))
    except Exception as e:
        logger.error(f"❌ Error processing PDF: {e}")
    return documents

def chunk_and_store(docs: List[Document], domain: str):
    if not docs:
        return
    chunker = SemanticChunker(embedding_function)
    chunks = chunker.split_documents(docs)
    vectorstore = Chroma(
        persist_directory=VECTOR_DIRS[domain],
        embedding_function=embedding_function,
        collection_name=domain
    )
    vectorstore.add_documents(chunks)
    logger.info(f"✅ Stored {len(chunks)} chunks for domain '{domain}'.")

def run_processing():
    logger.info("🚀 Bắt đầu xử lý toàn bộ tài liệu...")
    for domain, folder in DOC_DIRS.items():
        all_docs = []
        for file in os.listdir(folder):
            path = os.path.join(folder, file)
            if file.endswith(".docx"):
                all_docs.extend(extract_from_docx(path, domain))
            elif file.endswith(".pdf"):
                all_docs.extend(extract_from_pdf(path, domain))
        chunk_and_store(all_docs, domain)
    logger.info("🏁 Hoàn tất xử lý tất cả tài liệu.")

if __name__ == "__main__":
    run_processing()
